from docx2pdf import convert
from pypdf import PdfWriter, PdfReader, PdfMerger
import fitz
from docx.shared import Cm, Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
import re
from datetime import datetime
import locale

def get_file_info(input_path):
    # retrieve filename
    filename = input_path.split('/')[-1].split('.')[0]
    # retrieve company name
    pattern = r'^\d+|\bjpn\b|\bchi\b|\bger\b|\b\d+\s\w+\s\d+\b'   # can be replaced to extract company name from database
    company_name = re.sub(pattern, '', filename).strip()
    doc = Document(input_path)

    # Extract reference number and determine document id and date
    found_acq_ref = False
    for p in doc.paragraphs:
        if found_acq_ref:
            break
            
        if 'ACQ_REF' in p.text:
            ref = re.search(r"[\w\d]+(\/[\w\d]+)+", p.text).group(0)
            ref_parts = ref.split('/')
            doc_type = ref_parts[0]
            id = ref_parts[1]
            date = ref_parts[2]
            found_acq_ref = True

    filename_parts = filename.split(' ')

    # Detect language specified in file name
    if filename_parts[-1]=="JA" :
        language = filename_parts[-1]
        date = f"{date[:4]}年{date[4:6]}月{date[6:]}日"
    elif filename_parts[-1]=="DE":
        language = filename_parts[-1]
        locale.setlocale(locale.LC_TIME, 'de_DE.UTF-8')
        date_obj = datetime.strptime(date, "%Y%m%d")
        date = date_obj.strftime("%d %B %Y")
    elif filename_parts[-1]=="ZH":
        language = filename_parts[-1]
        date = f"{date[:4]}年{date[4:6]}月{date[6:]}日"
    else:
        language = "EN"
        locale.setlocale(locale.LC_TIME, 'en.UTF-8')
        date_obj = datetime.strptime(date, "%Y%m%d")
        date = date_obj.strftime("%d %B %Y")

    file_info = {
       "filename": filename,
       "company_name": company_name,
       "doc_type": doc_type,
       "id": id,
       "date": date,
       "language": language
    }

    return file_info



def extract_IS_cover_page(input_path,temporary_files_path):
    doc = Document(input_path)
    # IS Cover page specification
    margin = {"top": 6,          # in cm
            "bottom": 1.5,
            "left": 8,
            "right": 1.5}
    fontsize = 8.5
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith("#(1)"):
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.startswith("#(2))"):
            while len(doc.paragraphs)>i:
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(margin["top"])
        section.bottom_margin = Cm(margin["bottom"])
        section.left_margin = Cm(margin["left"])
        section.right_margin = Cm(margin["right"])

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            font = run.font
            font.size = Pt(fontsize)

    for section in doc.sections:
        header = section.header
        if header is not None:
            for p in header.paragraphs:
                p.text = ''
            footer = section.footer
        if footer is not None:
            for p in footer.paragraphs:
                p.text = ''

    doc.save(f'{temporary_files_path}/cover_page.docx')
    cover_page = f'{temporary_files_path}/cover_page.pdf'
    convert(f'{temporary_files_path}/cover_page.docx', cover_page)
    return cover_page

def extract_IS_content(input_path,temporary_files_path):
    doc = Document(input_path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.startswith("#(3)"):
            break
    for j in range(i):
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
    doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    doc.save(f'{temporary_files_path}/content.docx')
    content = f'{temporary_files_path}/content.pdf'
    convert(f'{temporary_files_path}/content.docx', f'{temporary_files_path}/content.pdf')
    return content

def create_IS_cover_page(cover_page, template_files_path, doc_type, language):
    template = PdfReader(open(f"{template_files_path}/cover-{doc_type}-{language}.pdf", "rb"))
    cover_content = PdfReader(open(cover_page, "rb"))
    for i in range(len(cover_content.pages)):
        background = template.pages[i]
        foreground = cover_content.pages[i]
        background.merge_page(foreground)

    writer = PdfWriter()
    number_of_cover = len(cover_content.pages)
    for i in range(number_of_cover):
        writer.add_page(template.pages[i])

    with open(cover_page, "wb") as outFile:
        writer.write(outFile)
    return number_of_cover

def merge_IS_pdf(cover_page, content, temporary_files_path):
    pdfs = [cover_page, content]
    merger = PdfMerger()
    for pdf in pdfs:
        merger.append(pdf)
    merged_pdf = f"{temporary_files_path}/output.pdf"
    merger.write(merged_pdf)
    merger.close()
    return merged_pdf

def edit_IS_pdf(merged_pdf, company_name, date, id, language, number_of_cover, filename, output_folder_path,fontfile):
    doc = fitz.open(merged_pdf) 
    number = f"NO.: {id}-{language}"
    page = doc[0]
    rect = fitz.Rect(225, 100, 650, 400) 
    rc = page.insert_textbox(rect, company_name, fontsize = 16, # choose fontsize (float)
                        fontname = "ArialUnicodeBold",       # a PDF standard font
                        fontfile = fontfile,
                        fill=(1, 0, 0),              
                        align = 0)                      # 0 = left, 1 = center, 2 = right

    rect1 = fitz.Rect(225, 130, 650, 400) 
    rc1 = page.insert_textbox(rect1, date, fontsize = 9, # choose fontsize (float)
                        fontname = "ArialUnicodeBold",       # a PDF standard font
                        fontfile = fontfile,
                        align = 0)                      # 0 = left, 1 = center, 2 = right

    rect2 = fitz.Rect(120, 80, 600, 400)
    rc1 = page.insert_textbox(rect2, number, fontsize = 9, # choose fontsize (float)
                        fontname = "ArialUnicodeBold",       # a PDF standard font
                        fontfile = fontfile,
                        align = 0)                      # 0 = left, 1 = center, 2 = right

    for i in range(1, number_of_cover):      # Assuming the information is added to all the cover pages
        page = doc[i]
        rect = fitz.Rect(225, 100, 650, 400) 
        rc = page.insert_textbox(rect, company_name, fontsize = 16, # choose fontsize (float)
                            fontname = "ArialUnicodeBold",       # a PDF standard font
                            fontfile = fontfile,
                            fill=(1, 0, 0),              
                            align = 0)                      # 0 = left, 1 = center, 2 = right

        rect1 = fitz.Rect(225, 130, 650, 400) 
        rc1 = page.insert_textbox(rect1, date, fontsize = 9, # choose fontsize (float)
                            fontname = "ArialUnicodeBold",       # a PDF standard font
                            fontfile = fontfile,
                            align = 0)                      # 0 = left, 1 = center, 2 = right

        rect2 = fitz.Rect(100, 120, 600, 400)
        rc1 = page.insert_textbox(rect2, number, fontsize = 9, # choose fontsize (float)
                            fontname = "ArialUnicodeBold",       # a PDF standard font
                            fontfile = fontfile,
                            align = 0)                      # 0 = left, 1 = center, 2 = right

    doc.save(f"{output_folder_path}/{filename}.pdf")   # Update file by doc.saveIncr(). Save to new instead by doc.save("new.pdf",...)