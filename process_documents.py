import deepl, fitz, re, os, locale, time, shutil
from docx2pdf import convert
from pypdf import PdfWriter, PdfReader
from docx.shared import Cm, Pt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from datetime import datetime

config = {
    "auth_key" : "daa69283-7136-c9a5-c283-38cbd645e38e",
    "input_folder" : "./docx_folder",
    "output_folder" : "./pdf_folder",
    "temporary_files_path": "./lib/temporary_files",
    "template_files_path": "./lib/templates",
    "font_file": "./lib/font/ArialUnicodeBold.ttf",
    "font_name": "ArialUnicodeBold",
    "CS_name_position": {"topleft_x": 90, 
                         "topleft_y": 265, 
                         "bottomright_x": 550, 
                         "bottomright_y": 400},
    "CS_name_fontsize": 28,
    "CS_date_position": {"topleft_x": 90, 
                         "topleft_y": 320, 
                         "bottomright_x": 650, 
                         "bottomright_y": 400},
    "CS_date_fontsize": 12,
    "CS_id_position": {"topleft_x": 90, 
                       "topleft_y": 110, 
                       "bottomright_x": 600, 
                       "bottomright_y": 400},
    "CS_id_fontsize": 9,  
    "IS_name_position": {"topleft_x": 225, 
                         "topleft_y": 100, 
                         "bottomright_x": 630, 
                         "bottomright_y": 400},
    "IS_name_fontsize": 16,
    "IS_date_position": {"topleft_x": 225, 
                         "topleft_y": 130, 
                         "bottomright_x": 650, 
                         "bottomright_y": 400},
    "IS_date_fontsize": 9,
    "IS_id_position": {"topleft_x": 120, 
                       "topleft_y": 80, 
                       "bottomright_x": 600, 
                       "bottomright_y": 400},
    "IS_id_fontsize": 9,  
    "AM_EM_date_position": {"topleft_x": 500, 
                         "topleft_y": 125, 
                         "bottomright_x": 650, 
                         "bottomright_y": 400},
    "AM_EM_date_fontsize": 9,
    "AM_EM_id_position": {"topleft_x": 120, 
                       "topleft_y": 80, 
                       "bottomright_x": 600, 
                       "bottomright_y": 400},
    "AM_EM_id_fontsize": 9,  
    "IR_name_position": {"topleft_x": 95, 
                         "topleft_y": 220, 
                         "bottomright_x": 630, 
                         "bottomright_y": 600},
    "IR_name_fontsize": 48,
    "IR_date_position": {"topleft_x": 95, 
                         "topleft_y": 300, 
                         "bottomright_x": 650, 
                         "bottomright_y": 500},
    "IR_date_fontsize": 12,
    "IR_id_position": {"topleft_x": 95, 
                       "topleft_y": 100, 
                       "bottomright_x": 600, 
                       "bottomright_y": 400},
    "IR_id_fontsize": 12,  
    "IS_EM_AM_margin": {"top": 5.5, 
                        "bottom": 1.5, 
                        "left": 8, 
                        "right": 1.5},  # For cover page
    "IS_EM_AM_text_fontsize": 8.5, # cover page
    "IR_margin": {"top": 17, 
                  "bottom": 1.5, 
                  "left": 11, 
                  "right": 1.5},  # For cover page
    "IR_text_fontsize": 9, # cover page
    
}            

def delete_files_in_temporary_folder():
    # Delete temporary files to avoid unexpected results

    folder_path = config["temporary_files_path"]
    max_attempts = 5  # Maximum number of attempts to delete a file
    attempt_delay = 1  # Delay in seconds between attempts
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        
        # Check if it's a file or symlink
        if os.path.isfile(file_path) or os.path.islink(file_path):
            attempts = 0
            while attempts < max_attempts:
                try:
                    os.remove(file_path)
                    break  # If we got here, the file was removed successfully
                except PermissionError:
                    time.sleep(attempt_delay)
                except Exception as e:
                    print(f"An error occurred while deleting {file_path}: {e}")
                    break  # Break the loop to avoid an infinite loop
                attempts += 1
            else:
                print(f"Failed to delete {file_path} after {max_attempts} attempts.")
                try:
                    shutil.rmtree(config["temporary_files_path"])
                    os.mkdir(config["temporary_files_path"])
                except Exception as e:
                    print("Failed to delete folder.")

        elif os.path.isdir(file_path):
            try:
                shutil.rmtree(file_path)
            except Exception as e:
                print(f"An error occurred while deleting {file_path}: {e}")

def format_date_A(date, language):
    # Format date to year, month, and date
    if language =="JA" or language =="ZH":
        date = f"{date[:4]}年{date[4:6]}月{date[6:]}日"

    elif language =="DE":
        locale.setlocale(locale.LC_TIME, 'de_DE.UTF-8')
        date_obj = datetime.strptime(date, "%Y%m%d")
        date = date_obj.strftime("%d %B %Y")

    elif language == "EN":
        locale.setlocale(locale.LC_TIME, 'en.UTF-8')
        date_obj = datetime.strptime(date, "%Y%m%d")
        date = date_obj.strftime("%d %B %Y")
    
    return date

def format_date_B(date, language):
    # Format date to year and month
    if language =="JA" or language =="ZH":
        date = f"{date[:4]}年{date[4:6]}月"

    elif language =="DE":
        locale.setlocale(locale.LC_TIME, 'de_DE.UTF-8')
        date_obj = datetime.strptime(date, "%Y%m%d")
        date = date_obj.strftime("%B %Y")

    elif language == "EN":
        locale.setlocale(locale.LC_TIME, 'en.UTF-8')
        date_obj = datetime.strptime(date, "%Y%m%d")
        date = date_obj.strftime("%B %Y")
    
    return date

def get_file_info(input_path):
    try:
        filename = input_path.split('/')[-1].split('\\')[-1].split('.')[0]

        file_info = {}

        # Retrieve company or industry name
        pattern = r'^\d+|\bJA\b|\bZH\b|\bDE\b|\b\d+\s\w+\s\d+\b'   
        name = re.sub(pattern, '', filename).strip()  # Can be replaced to extract company name from database
        with open(input_path, 'rb') as f:
            doc = Document(f)
 
        # Extract ACQ reference number and determine document id and date
        found_acq_ref = False
        for p in doc.paragraphs:
            if found_acq_ref:
                break

            if 'ACQ_REF' in p.text:
                ref = re.search(r"[\w\d]+\s*(\/\s*[\w\d]+\s*)+", p.text).group(0)
                ref_parts = ref.split('/')
                doc_type = ref_parts[0]
                id = ref_parts[1]
                date = ref_parts[2]
                if len(ref_parts) > 3:
                    database_code = ref_parts[3]
                    file_info["database_code"] = database_code
                found_acq_ref = True

        if not found_acq_ref:
            print(f"ACQ_REF not found for {filename}")
            return

        # Detect language specified in file name
        filename_parts = filename.split(' ')
        if filename_parts[-1]=="JA" or filename_parts[-1]=="ZH" or filename_parts[-1]=="DE":
            language = filename_parts[-1]
        else:
            language = "EN"
        
        # Format date according to document type
        if doc_type == "IS" or doc_type == "CS":
            date = format_date_A(date, language)
        elif doc_type == "AM" or doc_type == "EM" or doc_type == "IR":
            date = format_date_B(date, language)
        else:
            print("Wrong document type in ACQ_REF")
            return
        
        if doc_type == "IS":
            file_info["industry_name"] = name
        elif doc_type == "CS":
            file_info["company_name"] = name
        elif doc_type == "IR":
            found_acq_title = False
            for p in doc.paragraphs:
                if found_acq_title:
                    break

                if 'ACQ_TITLE' in p.text:
                    industry_name = p.text.split(": ")[-1]
                    found_acq_title = True

            if not found_acq_title:
                print(f"ACQ_TITLE not found for {filename}")
                return
            
            file_info["industry_name"] = industry_name
        else:
            pass

        file_info["filename"] = filename
        file_info["doc_type"] = doc_type
        file_info["id"] = id
        file_info["date"] = date
        file_info["language"] = language

        return file_info
    
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while getting file info.")
        return 


'''
# Alternative way to extract company name
from pymongo import MongoClient
def extract_company_name(company_code):

    try:
                # replace xxxxx with database ip address 
        client = MongoClient("mongodb://xxxxx:27017/") 
        # replace xxxxx with database name 
        db = client['xxxxx'] 
        # replace xxxxx with table wanted to used
        col = db['xxxxx'] 
        print("Connect to DB.")
    except:
        print("Can't connect to DB.")    

    documents = col.find({"compnumber": 111754})

    for document in documents:
        print(document['name'])

'''


def delete_text(input_path, list_of_words):
    # Delete ACQ_REF, ACQ_AUTHOR and ACQ_TITLE
    try:
        with open(input_path, 'rb') as f:
            doc = Document(f)

        for word in list_of_words:
            found_successfully = False
            for p in doc.paragraphs:
                if found_successfully:
                    break

                if word in p.text:
                    p.text = ""
                    found_successfully = True

            if not found_successfully:
                print(f"{word} not found in document")
                return

        file_path = os.path.join(config['temporary_files_path'], "input.docx")
        with open(file_path, 'wb') as f:
            doc.save(f)
        
        return file_path
    
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while deleting ACQ_REF, ACQ_AUTHOR or ACQ_TITLE.")
        return


def create_CS_pdf(docx_file, doc_type, language):
    # Merge CS content with the cover page template
    try:
        content_path = os.path.join(config['temporary_files_path'], "content.pdf")
        convert(docx_file, content_path)
        
        template = os.path.join(config['template_files_path'], f"cover-{doc_type}-{language}.pdf")
        
        with fitz.open(template) as cover_page:
            with fitz.open(content_path) as content:
                cover_page.insert_file(content)

                merged_pdf = os.path.join(config['temporary_files_path'], "output.pdf")
                cover_page.save(merged_pdf)

                return merged_pdf
    
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while creating the CS PDF.")
        return 

def edit_CS_pdf(merged_pdf, company_name, date, id, language, filename, output_folder):
    # Add company name and other information onto the cover page
    try:
        with fitz.open(merged_pdf) as doc:
            number = f"NO.: {id}-{language}"
            page = doc[0]
            
            rect_name = fitz.Rect(config["CS_name_position"]["topleft_x"], config["CS_name_position"]["topleft_y"], 
                             config["CS_name_position"]["bottomright_x"], config["CS_name_position"]["bottomright_y"])
            page.insert_textbox(rect_name, company_name, fontsize=config["CS_name_fontsize"], fontname=config["font_name"],
                                    fontfile=config["font_file"], fill=(1, 0, 0), align=0)

            rect_date = fitz.Rect(config["CS_date_position"]["topleft_x"], config["CS_date_position"]["topleft_y"],
                              config["CS_date_position"]["bottomright_x"], config["CS_date_position"]["bottomright_y"])
            page.insert_textbox(rect_date, date, fontsize=config["CS_date_fontsize"], fontname=config["font_name"],
                                    fontfile=config["font_file"], align=0)

            rect_id = fitz.Rect(config["CS_id_position"]["topleft_x"], config["CS_id_position"]["topleft_y"],
                              config["CS_id_position"]["bottomright_x"], config["CS_id_position"]["bottomright_y"])
            page.insert_textbox(rect_id, number, fontsize=config["CS_id_fontsize"], fontname=config["font_name"],
                                    fontfile=config["font_file"], align=0)

            doc.save(os.path.join(output_folder, f"{filename}.pdf"))
        
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while editing the CS PDF.")

def extract_IS_EM_AM_cover_page(input_path):
    # Extract cover page and set to correct margin
    try:
        with open(input_path, 'rb') as f:
            doc = Document(f)
        # Delete table content
        for table in doc.tables:
            table._element.getparent().remove(table._element)
        # Delete #(1) line
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith("#(1)"):
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
                break
        # Delete everything behind #(2) line
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith("#(2)"):
                while len(doc.paragraphs) > i:
                    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        # Set to correct margin and delete header and footer
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(config["IS_EM_AM_margin"]["top"])
            section.bottom_margin = Cm(config["IS_EM_AM_margin"]["bottom"])
            section.left_margin = Cm(config["IS_EM_AM_margin"]["left"])
            section.right_margin = Cm(config["IS_EM_AM_margin"]["right"])
            
            header = section.header
            if header is not None:
                for p in header.paragraphs:
                    p.text = ''
            footer = section.footer
            if footer is not None:
                for p in footer.paragraphs:
                    p.text = ''
        # Format paragraph spacing, alignment and fontsize
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(config["IS_EM_AM_text_fontsize"])
        # Save the document
        docx_path = os.path.join(config['temporary_files_path'], "cover_page.docx")
        with open(docx_path, 'wb') as f:
            doc.save(f)
        # Convert to PDF
        cover_page = os.path.join(config['temporary_files_path'], "cover_page.pdf")
        convert(docx_path, cover_page)

        return cover_page
    
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while extracting IS/EM/AM cover page.")
        return

def extract_IR_cover_page(input_path):
    # Extract cover page and set to correct margin
    try:
        with open(input_path, 'rb') as f:
            doc = Document(f)
        # Delete table content
        for table in doc.tables:
            table._element.getparent().remove(table._element)
        # Delete #(1) line
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith("#(1)"):
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
                break
        # Delete everything behind #(2) line
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.startswith("#(2)"):
                while len(doc.paragraphs) > i:
                    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        # Set to correct margin and delete header and footer
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(config["IR_margin"]["top"])
            section.bottom_margin = Cm(config["IR_margin"]["bottom"])
            section.left_margin = Cm(config["IR_margin"]["left"])
            section.right_margin = Cm(config["IR_margin"]["right"])
            
            header = section.header
            if header is not None:
                for p in header.paragraphs:
                    p.text = ''
            footer = section.footer
            if footer is not None:
                for p in footer.paragraphs:
                    p.text = ''
        # Format paragraph spacing, alignment and fontsize
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(config["IR_text_fontsize"])
        # Save the document
        docx_path = os.path.join(config['temporary_files_path'], "cover_page.docx")
        with open(docx_path, 'wb') as f:
            doc.save(f)
        # Convert to PDF
        cover_page = os.path.join(config['temporary_files_path'], "cover_page.pdf")
        convert(docx_path, cover_page)

        return cover_page
    
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while extracting IR cover page.")
        return

def extract_IS_EM_AM_IR_content(input_path):
    # Extract IS/EM/AM/IR content
    try:
        with open(input_path, 'rb') as f:
            doc = Document(f)
        # Delete everything before #(3)
        while not doc.paragraphs[0].text.startswith("#(3)"):
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
        # Delete the line #(3)     
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
        # Save the content
        docx_path = os.path.join(config['temporary_files_path'],"content.docx")
        with open(docx_path, 'wb') as f:
            doc.save(f)
        # Convert docx to pdf
        content = os.path.join(config['temporary_files_path'], "content.pdf")
        convert(docx_path, content)
        
        return content
    
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while extracting IS content.")
        return 

def overlay_cover_page(template, cover_content, cover_page):
    # Overlay conver page content
    for i in range(len(cover_content.pages)):
        background = template.pages[i]
        foreground = cover_content.pages[i]
        background.merge_page(foreground)

    writer = PdfWriter()
    number_of_cover = len(cover_content.pages)
    for i in range(number_of_cover):
        writer.add_page(template.pages[i])

    with open(cover_page, "wb") as outFile:
        writer.write(outFile)
    return number_of_cover

def create_IS_EM_AM_IR_cover_page(cover_page, doc_type, language):
    try:
        template_path = os.path.join(config['template_files_path'],f"cover-{doc_type}-{language}.pdf")
        # Read template and cover page content pdf
        with open(template_path, "rb") as f:
            template = PdfReader(f)
            with open(cover_page, "rb") as g:
                cover_content = PdfReader(g)
                # Handle the case where the cover page content pdf has more page than the template
                if len(cover_content.pages) > len(template.pages):
                    with fitz.open(template_path) as t1:
                        with fitz.open(template_path) as t2:
                            # Extract the last page of the template
                            for _ in range(len(template.pages)-1):
                                t2.delete_page(0)
                            # Append additional page to the template
                            for _ in range(len(cover_content.pages) - len(template.pages)):
                                t1.insert_file(t2)
                            # Save the file
                            temp = os.path.join(config['temporary_files_path'],"template.pdf")
                            t1.save(temp)
                            # Load the file and overlay the content
                            with open(temp, "rb") as h:
                                template = PdfReader(h)
                                number_of_cover = overlay_cover_page(template, cover_content, cover_page)
                else: 
                    pass
                # Overlay the content
                number_of_cover = overlay_cover_page(template, cover_content, cover_page)
                return number_of_cover

    except Exception as e:
        print(f"{str(e)}\nAn error occurred while creating IS/EM/AM/IR cover page.")
        return 

def merge_pdf(cover_page, content):
    try:
        with fitz.open(cover_page) as merger1:
            with fitz.open(content) as merger2:
                # Merge pdf
                merger1.insert_file(merger2)
                # Save the merged pdf
                merged_pdf = os.path.join(config['temporary_files_path'], "output.pdf")
                merger1.save(merged_pdf)

                return merged_pdf
            
    except Exception as e:
        print(f"{str(e)}\nAn error occurred while merging IS PDFs.")
        return 

def edit_IS_pdf(merged_pdf, industry_name, date, id, language, number_of_cover, filename, output_folder):
    try:
        with fitz.open(merged_pdf) as doc:
            # Translate industry name if it is not in English
            if language != "EN":
                translator = deepl.Translator(config["auth_key"])
                result = translator.translate_text(industry_name, target_lang=language)
                industry_name = result.text
            # Add industry name
            def add_industry_name(page):
                rect = fitz.Rect(config["IS_name_position"]["topleft_x"], config["IS_name_position"]["topleft_y"], 
                                 config["IS_name_position"]["bottomright_x"], config["IS_name_position"]["bottomright_y"])
                page.insert_textbox(rect, industry_name, fontsize=config["IS_name_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], fill=(1, 0, 0), align=0)
            # Add date
            def add_date(page):
                rect = fitz.Rect(config["IS_date_position"]["topleft_x"], config["IS_date_position"]["topleft_y"], 
                                 config["IS_date_position"]["bottomright_x"], config["IS_date_position"]["bottomright_y"])
                page.insert_textbox(rect, date, fontsize=config["IS_date_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add id
            number = f"NO.: {id}-{language}"
            def add_id(page):
                rect = fitz.Rect(config["IS_id_position"]["topleft_x"], config["IS_id_position"]["topleft_y"], 
                                 config["IS_id_position"]["bottomright_x"], config["IS_id_position"]["bottomright_y"])
                page.insert_textbox(rect, number, fontsize=config["IS_id_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add everthing onto the first page
            page = doc[0]
            add_industry_name(page)
            add_date(page)
            add_id(page)
            # Assuming the information is added to all the cover pages
            for i in range(1, number_of_cover):  
                page = doc[i]
                add_industry_name(page)
                add_date(page)
                add_id(page)
            # Save the file to the output folder
            doc.save(os.path.join(output_folder, f"{filename}.pdf"))

    except Exception as e:
        print(f"{str(e)}\nAn error occurred while editing IS PDF.")

def edit_AM_EM_pdf(merged_pdf, date, id, language, number_of_cover, filename, output_folder):
    try:
        with fitz.open(merged_pdf) as doc:
            # Add date
            def add_date(page):
                rect = fitz.Rect(config["AM_EM_date_position"]["topleft_x"], config["AM_EM_date_position"]["topleft_y"], 
                                 config["AM_EM_date_position"]["bottomright_x"], config["AM_EM_date_position"]["bottomright_y"])
                page.insert_textbox(rect, date, fontsize=config["AM_EM_date_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add id
            number = f"NO.: {id}-{language}"
            def add_id(page):
                rect = fitz.Rect(config["AM_EM_id_position"]["topleft_x"], config["AM_EM_id_position"]["topleft_y"], 
                                 config["AM_EM_id_position"]["bottomright_x"], config["AM_EM_id_position"]["bottomright_y"])
                page.insert_textbox(rect, number, fontsize=config["AM_EM_id_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add the information onto the first page
            page = doc[0]
            add_date(page)
            add_id(page)
            # Assuming the information is added to all the cover pages
            for i in range(1, number_of_cover):  
                page = doc[i]
                add_date(page)
                add_id(page)
            # Save the file to the output folder
            doc.save(os.path.join(output_folder, f"{filename}.pdf"))

    except Exception as e:
        print(f"{str(e)}\nAn error occurred while editing AM/EM PDF.")

def edit_IR_pdf(merged_pdf, industry_name, date, id, language, number_of_cover, filename, output_folder):
    try:
        with fitz.open(merged_pdf) as doc:
            # Translate industry name if it is not in English
            if language != "EN":
                translator = deepl.Translator(config["auth_key"])
                result = translator.translate_text(industry_name, target_lang=language)
                industry_name = result.text
            # Add industry name
            def add_industry_name(page):
                rect = fitz.Rect(config["IR_name_position"]["topleft_x"], config["IR_name_position"]["topleft_y"], 
                                 config["IR_name_position"]["bottomright_x"], config["IR_name_position"]["bottomright_y"])
                page.insert_textbox(rect, industry_name, fontsize=config["IR_name_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add date
            def add_date(page):
                rect = fitz.Rect(config["IR_date_position"]["topleft_x"], config["IR_date_position"]["topleft_y"], 
                                 config["IR_date_position"]["bottomright_x"], config["IR_date_position"]["bottomright_y"])
                page.insert_textbox(rect, date, fontsize=config["IR_date_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add id
            number = f"NO.: {id}-{language}"
            def add_id(page):
                rect = fitz.Rect(config["IR_id_position"]["topleft_x"], config["IR_id_position"]["topleft_y"], 
                                 config["IR_id_position"]["bottomright_x"], config["IR_id_position"]["bottomright_y"])
                page.insert_textbox(rect, number, fontsize=config["IR_id_fontsize"], fontname=config["font_name"],
                                        fontfile=config["font_file"], align=0)
            # Add everthing onto the first page
            page = doc[0]
            add_industry_name(page)
            add_date(page)
            add_id(page)
            # Assuming the information is added to all the cover pages
            for i in range(1, number_of_cover):  
                page = doc[i]
                add_industry_name(page)
                add_date(page)
                add_id(page)
            # Save the file to the output folder
            doc.save(os.path.join(output_folder, f"{filename}.pdf"))

    except Exception as e:
        print(f"{str(e)}\nAn error occurred while editing IR PDF.")

def main():
    # Load file in docx folder
    translated_files = []
    for file in os.listdir(config["input_folder"]):
        translated_files.append(os.path.join(config["input_folder"], file))
    
    for file in translated_files:
        print(f"Now processing {file}")
            
        try:
            file_info = get_file_info(file)
            filename = file_info["filename"]
            doc_type = file_info["doc_type"]
            id = file_info["id"]
            date = file_info["date"]
            language = file_info["language"]

            if doc_type == "IS":
                industry_name = file_info["industry_name"]
                file_path = delete_text(file, ["ACQ_REF", "ACQ_AUTHOR"])
                cover_page = extract_IS_EM_AM_cover_page(file_path)
                content = extract_IS_EM_AM_IR_content(file_path)
                number_of_cover = create_IS_EM_AM_IR_cover_page(cover_page, doc_type, language) 
                merged_pdf = merge_pdf(cover_page, content)
                edit_IS_pdf(merged_pdf, industry_name, date, id, language, number_of_cover, filename, config["output_folder"])
            elif doc_type == "CS":
                company_name = file_info["company_name"]
                file_path = delete_text(file, ["ACQ_REF", "ACQ_AUTHOR"])
                merged_pdf = create_CS_pdf(file_path, doc_type, language)
                edit_CS_pdf(merged_pdf, company_name, date, id, language, filename,config["output_folder"])
            elif doc_type == "AM" or doc_type == "EM":
                file_path = delete_text(file, ["ACQ_REF", "ACQ_AUTHOR"])
                cover_page = extract_IS_EM_AM_cover_page(file_path)
                content = extract_IS_EM_AM_IR_content(file_path)
                number_of_cover = create_IS_EM_AM_IR_cover_page(cover_page, doc_type, language)
                merged_pdf = merge_pdf(cover_page, content) 
                edit_AM_EM_pdf(merged_pdf, date, id, language, number_of_cover, filename, config["output_folder"])
            
            elif doc_type == "IR":
                industry_name = file_info["industry_name"]
                file_path = delete_text(file, ["ACQ_REF", "ACQ_AUTHOR", "ACQ_TITLE"])
                cover_page = extract_IR_cover_page(file_path)
                content = extract_IS_EM_AM_IR_content(file_path)
                number_of_cover = create_IS_EM_AM_IR_cover_page(cover_page, doc_type, language)
                merged_pdf = merge_pdf(cover_page, content) 
                edit_IR_pdf(merged_pdf,industry_name, date, id, language, number_of_cover, filename, config["output_folder"])

        except Exception as e:
            print(f"{e}\nAn error occurs when processing {file}")
            continue

        delete_files_in_temporary_folder()

if __name__ == "__main__":
    main()